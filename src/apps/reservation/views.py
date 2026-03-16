import os
from django.http import HttpResponse
from pathlib import Path
from django.conf import settings

import calendar
from datetime import datetime, time, timedelta

from django.db import transaction
from django.db.models import Q

from rest_framework.views import APIView
from rest_framework.decorators import action
from rest_framework import generics, viewsets, serializers, status
from rest_framework.response import Response
from rest_framework.exceptions import ValidationError
from rest_framework.permissions import AllowAny, IsAuthenticated

from drf_spectacular.utils import OpenApiParameter, OpenApiTypes, extend_schema
from .serializers import ReservationSerializer, ReservationListSerializer, ReservationRetrieveSerializer, ClientReservationSerializer, CalendarReservationSerializer, ReciptSerializer

from apps.core.paginator import CustomPagination
from slugify import slugify

from .models import Reservation, RentalReceipt, Clients, Property

from apps.accounts.models import CustomUser

from apps.core.functions import get_month_name, generate_audit, check_user_has_rol, confeccion_ics
from apps.dashboard.utils import get_stadistics_period
import subprocess
from docxtpl import DocxTemplate
from babel.dates import format_date
import io
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from django.utils.timezone import now
from .signals import send_purchase_event_to_meta
from django.views.decorators.csrf import csrf_exempt
import json
from django.http import JsonResponse, HttpResponseBadRequest

class ReservationsApiView(viewsets.ModelViewSet):
    serializer_class = ReservationSerializer
    queryset = Reservation.objects.exclude(deleted=True).order_by("check_in_date")
    permission_classes = [IsAuthenticated]  # Requiere autenticación
    search_fields = [
        "client__email",
        "client__first_name",
        "client__last_name",
        "property__name"
    ]
    pagination_class = CustomPagination

    def get_pagination_class(self):
        """Determinar si usar o no paginación
        - page_size = valor
        - valor = un numero entero, será el tamaño de la pagina
        - valor = none, no se pagina el resultado
        """
        if self.request.GET.get("page_size") == "none":
            return None

        return self.pagination_class

    def get_queryset(self):
        queryset = super().get_queryset()

        """
        Custom queryset to search reservations in a given month-year
        """
        if self.action == 'list':
            if self.request.query_params:
                from_param = self.request.query_params.get('from')
                type_param = self.request.query_params.get('type')
                now = datetime.now()
                check_in_time = time(15, 0)  # 3 PM
                check_out_time = time(11, 0)  # 11 AM

                if self.request.query_params.get('year') and self.request.query_params.get('month'):
                    try:
                        month_param = int(self.request.query_params['month'])
                        if not month_param in range(1, 13):
                            raise ValidationError({"error": "Parámetro Mes debe ser un número entre el 1 y el 12"})

                    except Exception:
                        raise ValidationError({"error_month_param": "Parámetro Mes debe ser un número entre el 1 y el 12"})

                    try:
                        year_param = int(self.request.query_params['year'])
                        if year_param < 1:
                            raise ValidationError({"error": "Parámetro Año debe ser un número entero positivo"})

                    except Exception:
                        raise ValidationError({"error_year_param": "Año debe ser un número entero positivo"})

                    last_day_month = calendar.monthrange(year_param, month_param)[1]

                    range_evaluate = (datetime(year_param, month_param, 1), datetime(year_param, month_param, last_day_month))

                    if self.request.query_params.get('from_check_in') == 'true':
                        queryset = queryset.filter(check_in_date__range=range_evaluate)
                    else:
                        queryset = queryset.filter(
                            Q(check_in_date__range=range_evaluate) |
                            Q(check_out_date__range=range_evaluate)
                        )

                if from_param == 'today':
                    queryset = queryset.filter(check_in_date__gte=now)
                elif from_param == 'in_progress':
                    # Verificar reservas que están en curso
                    now_time = now.time()
                    today_date = now.date()
                    tomorrow_date = today_date + timedelta(days=1)

                    queryset = queryset.filter(
                        Q(check_in_date__lt=today_date, check_out_date__gt=today_date) |
                        (Q(check_in_date=today_date) & Q(check_out_date__gt=today_date)) |
                        (Q(check_out_date=today_date) & Q(check_out_date__gt=now)) |
                        (Q(check_out_date=tomorrow_date) & Q(check_out_date__gt=datetime.combine(tomorrow_date, check_out_time)))
                    )

                elif from_param == 'pending':
                    queryset = queryset.filter(status__in=['pending', 'under_review'])

                if self.request.query_params.get('type'):
                    queryset = queryset.filter(origin=self.request.query_params.get('type'))
                
                # Filtrar por reservas creadas hoy
                if self.request.query_params.get('created_today') == 'true':
                    from django.utils import timezone
                    # Obtener la fecha local de hoy
                    local_now = timezone.localtime()
                    today = local_now.date()
                    
                    # Crear el inicio y fin del día en la zona horaria local
                    start_of_day_local = timezone.make_aware(
                        datetime.combine(today, time.min),
                        timezone.get_current_timezone()
                    )
                    end_of_day_local = timezone.make_aware(
                        datetime.combine(today, time.max),
                        timezone.get_current_timezone()
                    )
                    
                    # Filtrar usando rango en UTC (evita problemas con CONVERT_TZ en MySQL)
                    queryset = queryset.filter(
                        created__gte=start_of_day_local,
                        created__lte=end_of_day_local
                    )

        elif self.action in ['partial_update', 'update', 'destroy']:
            if not check_user_has_rol("admin", self.request.user):
                queryset = queryset.filter(
                    Q(origin="air") |
                    Q(seller=self.request.user)
                )

        if self.request.query_params.get('exclude'):
            queryset = queryset.exclude(origin=self.request.query_params['exclude'])

        return queryset.exclude(deleted=True)

    def get_serializer_class(self):
        if self.action == 'retrieve':
            return ReservationRetrieveSerializer
        if self.action == 'list':
            return ReservationListSerializer

        return super().get_serializer_class()

    def get_serializer_context(self):
        context = super().get_serializer_context()
        if self.action == 'retrieve':
            context['retrieve'] = True
        return context

    @extend_schema(
        parameters=[
            OpenApiParameter(
                "year",
                OpenApiTypes.INT,
                required=False,
                description="Filter results by year",
                enum=[2024, 2023, 2022],
            ),
            OpenApiParameter(
                "month",
                OpenApiTypes.INT,
                required=False,
                description="Filter results by month (number 1 to 12)",
                enum=list(range(1, 13)),
            ),
            OpenApiParameter(
                "page_size",
                OpenApiTypes.INT,
                description="Enviar page_size=valor para determinar tamaño de la pagina, sino enviar page_size=none para no tener paginado",
                required=False,
            ),
            OpenApiParameter(
                "from",
                OpenApiTypes.STR,
                description="Obtiene todas las reservas desde la fecha de hoy en adelante. Se puede combinar con type (tipo de reserva), pero no con year o month",
                required=False,
                enum=["today", "in_progress", "pending"]
            ),
            OpenApiParameter(
                "from_check_in",
                OpenApiTypes.STR,
                description="Obtiene todas las reservas que hayan comenzado este mes. Se usa en combinacion de los queryparams year y month",
                required=False,
                enum=["true"]
            ),
            OpenApiParameter(
                "type",
                OpenApiTypes.STR,
                description="Filtra las resevas según donde se generaron, AriBnB (air), Sistema Casa Austin (aus), Mantenimiento (man)",
                required=False,
                enum=["aus", "air", "man"]
            ),
            OpenApiParameter(
                "exclude",
                OpenApiTypes.STR,
                description="Excluye las resevas según donde se generaron, AriBnB (air), Sistema Casa Austin (aus), Mantenimiento (man)",
                required=False,
                enum=["aus", "air", "man"]
            ),
            OpenApiParameter(
                name='id',
                description='A numeric ID identifying this reservation.',
                required=True,
                type=OpenApiTypes.INT,
                location=OpenApiParameter.PATH
            ),
        ],
        responses={200: ReservationListSerializer(many=True)},
        methods=["GET"],
    )
    def list(self, request, *args, **kwargs):
        self.pagination_class = self.get_pagination_class()
        return super().list(request, *args, **kwargs)

    def perform_create(self, serializer):
        with transaction.atomic():
            user_seller = self.request.user
            if self.request.POST['origin'].lower() == 'air':
                user_seller = CustomUser.objects.get(first_name='AirBnB')

            # Las reservas creadas por admin/vendedores están aprobadas por defecto
            instance = serializer.save(seller=user_seller, status='approved')
            if instance.late_checkout:
                original_check_out_date = instance.check_out_date
                instance.late_check_out_date = original_check_out_date
                instance.check_out_date = original_check_out_date + timedelta(days=1)
            else:
                instance.late_check_out_date = None
            instance.save()

            for file in self.request.FILES.getlist('file'):
                RentalReceipt.objects.create(
                    reservation=instance,
                    file=file
                )

        confeccion_ics()

        generate_audit(
            instance,
            self.request.user,
            "create",
            "Reserva creada"
        )

    def perform_update(self, serializer):
        with transaction.atomic():
            instance = self.get_object()
            original_late_checkout = instance.late_checkout
            original_check_out_date = instance.check_out_date
            instance = serializer.save()

            # Solo guardar si realmente hay cambios en late checkout
            if instance.late_checkout and not original_late_checkout:
                instance.late_check_out_date = original_check_out_date
                instance.check_out_date = original_check_out_date + timedelta(days=1)
                instance.save(update_fields=['late_check_out_date', 'check_out_date'])
            elif not instance.late_checkout and original_late_checkout:
                instance.check_out_date = instance.late_check_out_date
                instance.late_check_out_date = None
                instance.save(update_fields=['check_out_date', 'late_check_out_date'])

            confeccion_ics()

            generate_audit(
                instance,
                self.request.user,
                "update",
                "Reserva actualizada full"
            )

        return super().perform_update(serializer)

    def partial_update(self, request, *args, **kwargs):
        instance = self.get_object()

        serializer = self.get_serializer(instance, data=request.data, partial=True)
        serializer.is_valid(raise_exception=True)

        with transaction.atomic():
            original_late_checkout = instance.late_checkout
            original_check_out_date = instance.check_out_date

            # Determinar si hay cambio en late_checkout ANTES del save
            new_late_checkout = serializer.validated_data.get('late_checkout', instance.late_checkout)

            # Manejar cambios en late_checkout ANTES del save para evitar race conditions
            if new_late_checkout and not original_late_checkout:
                # Activando late_checkout: extender check_out_date +1 día
                serializer.validated_data['late_check_out_date'] = original_check_out_date
                serializer.validated_data['check_out_date'] = original_check_out_date + timedelta(days=1)
            elif not new_late_checkout and original_late_checkout:
                # Desactivando late_checkout: revertir check_out_date
                serializer.validated_data['check_out_date'] = instance.late_check_out_date
                serializer.validated_data['late_check_out_date'] = None

            instance = serializer.save()

            for file in request.FILES.getlist('file'):
                RentalReceipt.objects.create(
                    reservation=instance,
                    file=file
                )

        confeccion_ics()

        generate_audit(
            serializer.instance,
            self.request.user,
            "update",
            "Reserva actualizada"
        )
        return Response(serializer.data)


    def destroy(self, request, *args, **kwargs):
        instance = self.get_object()
        # Pasar el motivo de eliminación para auditoría con simple_history
        user_info = f"{request.user.email}" if request.user and hasattr(request.user, 'email') else "Usuario API"
        instance.delete(reason=f"API: eliminado por {user_info}")

        confeccion_ics()

        generate_audit(
            instance,
            self.request.user,
            "delete",
            "Reserva eliminada"
        )
        return Response(status=204)

###### MOD AUSTIN #######
    @action(detail=True, methods=['get'], url_path='contrato')
    def contrato(self, request, pk=None):
        try:
            reservation = self.get_object()
            client = Clients.objects.get(id=reservation.client_id)
            property = Property.objects.get(id=reservation.property_id)

            document_type_map = {
                'pas': 'pasaporte',
                'cex': 'carnet de extranjería',
                'dni': 'DNI',
                'ruc': 'RUC'
            }

            document_type = document_type_map.get(client.document_type, None)
            if document_type is None:
                raise ValueError(f"Tipo de documento desconocido: {client.document_type}")

            checkin_date = format_date(reservation.check_in_date, format="d 'de' MMMM 'del' YYYY", locale='es')
            checkout_date = format_date(reservation.check_out_date, format="d 'de' MMMM 'del' YYYY", locale='es')

            # Determinar qué plantilla usar según el tipo de documento
            if client.document_type == 'ruc':
                template_path = os.path.join(os.path.dirname(__file__), '../../plantilla_ruc.docx')
            else:
                template_path = os.path.join(os.path.dirname(__file__), '../../plantilla.docx')

            doc = DocxTemplate(template_path)

            context = {
                'nombre': f"{client.first_name.upper()} {client.last_name.upper()}",
                'tipodocumento': document_type.upper(),
                'dni': client.number_doc,
                'propiedad': property.name,
                'checkin': checkin_date,
                'checkout': checkout_date,
                'preciodolares': f"${reservation.price_usd:.2f}",
                'numpax': str(reservation.guests)
            }

            doc.render(context)

            temp_doc_path = "/tmp/temp_contract.docx"
            temp_pdf_path = "/tmp/temp_contract.pdf"
            doc.save(temp_doc_path)

            subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', '/tmp', temp_doc_path], check=True)

            with open(temp_pdf_path, "rb") as pdf_file:
                pdf_data = pdf_file.read()

            response = HttpResponse(pdf_data, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{property.name}_contract.pdf"'

            os.remove(temp_doc_path)
            os.remove(temp_pdf_path)

            return response
        except Clients.DoesNotExist:
            return Response({'error': 'Client not found'}, status=404)
        except Property.DoesNotExist:
            return Response({'error': 'Property not found'}, status=404)
        except ValueError as e:
            return Response({'error': str(e)}, status=400)
        except subprocess.CalledProcessError as e:
            return Response({'error': 'Error converting DOCX to PDF'}, status=500)
        except Exception as e:
            return Response({'error': str(e)}, status=400)
###### FIN MOD #######

    @action(detail=False, methods=['get'], url_path='contratos-zip')
    def contratos_zip(self, request):
        """Genera un ZIP con todos los contratos PDF de un mes dado."""
        import zipfile
        import tempfile
        import uuid

        month = request.query_params.get('month')  # formato: YYYY-MM
        if not month:
            return Response({'error': 'Parámetro month requerido (YYYY-MM)'}, status=400)

        try:
            year, mon = month.split('-')
            year, mon = int(year), int(mon)
            _, last_day = calendar.monthrange(year, mon)
        except (ValueError, TypeError):
            return Response({'error': 'Formato de month inválido. Usar YYYY-MM'}, status=400)

        from datetime import date
        date_from = date(year, mon, 1)
        date_to = date(year, mon, last_day)

        reservations = Reservation.objects.filter(
            check_in_date__gte=date_from,
            check_in_date__lte=date_to,
            status='approved',
            deleted=False,
            client__isnull=False,
        ).select_related('client', 'property').order_by('check_in_date')

        if not reservations.exists():
            return Response({'error': 'No hay reservas aprobadas en ese mes'}, status=404)

        document_type_map = {
            'pas': 'pasaporte',
            'cex': 'carnet de extranjería',
            'dni': 'DNI',
            'ruc': 'RUC',
        }

        # Crear directorio temporal único para este request
        tmp_dir = os.path.join(tempfile.gettempdir(), f'contracts_{uuid.uuid4().hex[:8]}')
        os.makedirs(tmp_dir, exist_ok=True)

        pdf_files = []
        errors = []

        try:
            for res in reservations:
                client = res.client
                prop = res.property

                doc_type = document_type_map.get(client.document_type)
                if not doc_type:
                    errors.append(f'{client.first_name}: tipo documento desconocido')
                    continue

                checkin_date = format_date(res.check_in_date, format="d 'de' MMMM 'del' YYYY", locale='es')
                checkout_date = format_date(res.check_out_date, format="d 'de' MMMM 'del' YYYY", locale='es')

                if client.document_type == 'ruc':
                    template_path = os.path.join(os.path.dirname(__file__), '../../plantilla_ruc.docx')
                else:
                    template_path = os.path.join(os.path.dirname(__file__), '../../plantilla.docx')

                try:
                    doc = DocxTemplate(template_path)
                    context = {
                        'nombre': f"{client.first_name.upper()} {(client.last_name or '').upper()}",
                        'tipodocumento': doc_type.upper(),
                        'dni': client.number_doc,
                        'propiedad': prop.name,
                        'checkin': checkin_date,
                        'checkout': checkout_date,
                        'preciodolares': f"${res.price_usd:.2f}",
                        'numpax': str(res.guests),
                    }
                    doc.render(context)

                    safe_name = slugify(f"{client.first_name}_{client.last_name or ''}")
                    file_base = f"{res.check_in_date.strftime('%Y-%m-%d')}_{safe_name}_{prop.name}"
                    docx_path = os.path.join(tmp_dir, f'{file_base}.docx')
                    pdf_path = os.path.join(tmp_dir, f'{file_base}.pdf')

                    doc.save(docx_path)
                    subprocess.run(
                        ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', tmp_dir, docx_path],
                        check=True, capture_output=True, timeout=30,
                    )

                    if os.path.exists(pdf_path):
                        pdf_files.append((f'{file_base}.pdf', pdf_path))
                    os.remove(docx_path)

                except Exception as e:
                    errors.append(f'{client.first_name} {res.check_in_date}: {str(e)}')
                    continue

            if not pdf_files:
                return Response({'error': 'No se pudo generar ningún contrato', 'details': errors}, status=500)

            # Crear ZIP en memoria
            zip_path = os.path.join(tmp_dir, 'contratos.zip')
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
                for filename, filepath in pdf_files:
                    zf.write(filepath, filename)

            with open(zip_path, 'rb') as f:
                zip_data = f.read()

            month_name = format_date(date_from, format='MMMM_YYYY', locale='es')
            response = HttpResponse(zip_data, content_type='application/zip')
            response['Content-Disposition'] = f'attachment; filename="contratos_{month_name}.zip"'
            return response

        finally:
            # Limpiar archivos temporales
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)

    @action(detail=True, methods=['get'], url_path='contrato-firma')
    def contrato_firma(self, request, pk=None):
        """Genera contrato PDF con la firma digitalizada del cliente (RENIEC)."""
        try:
            reservation = self.get_object()
            client = Clients.objects.get(id=reservation.client_id)
            property_obj = Property.objects.get(id=reservation.property_id)

            document_type_map = {
                'pas': 'pasaporte',
                'cex': 'carnet de extranjería',
                'dni': 'DNI',
                'ruc': 'RUC',
            }
            document_type = document_type_map.get(client.document_type)
            if document_type is None:
                raise ValueError(f"Tipo de documento desconocido: {client.document_type}")

            checkin_date = format_date(reservation.check_in_date, format="d 'de' MMMM 'del' YYYY", locale='es')
            checkout_date = format_date(reservation.check_out_date, format="d 'de' MMMM 'del' YYYY", locale='es')

            # Plantilla con placeholder {{ firma }}
            if client.document_type == 'ruc':
                template_path = os.path.join(os.path.dirname(__file__), '../../plantilla_ruc_firma.docx')
            else:
                template_path = os.path.join(os.path.dirname(__file__), '../../plantilla_firma.docx')

            doc = DocxTemplate(template_path)

            # --- Obtener firma del cliente (PNG bytes) ---
            firma_bytes = self._get_firma_bytes(client)

            context = {
                'nombre': f"{client.first_name.upper()} {(client.last_name or '').upper()}",
                'tipodocumento': document_type.upper(),
                'dni': client.number_doc,
                'propiedad': property_obj.name,
                'checkin': checkin_date,
                'checkout': checkout_date,
                'preciodolares': f"${reservation.price_usd:.2f}",
                'numpax': str(reservation.guests),
                'firma': '',  # Se inserta como imagen flotante después
            }

            doc.render(context)

            import uuid as _uuid
            uid = _uuid.uuid4().hex[:8]
            temp_doc_path = f"/tmp/contract_firma_{uid}.docx"
            temp_pdf_path = f"/tmp/contract_firma_{uid}.pdf"
            doc.save(temp_doc_path)

            # Insertar firma como imagen flotante (anchor) encima de la línea
            if firma_bytes:
                self._insert_firma_anchor(temp_doc_path, firma_bytes)

            subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', '/tmp', temp_doc_path],
                check=True,
            )

            with open(temp_pdf_path, "rb") as pdf_file:
                pdf_data = pdf_file.read()

            # Aplicar efecto de documento escaneado (determinístico por cliente)
            pdf_data = self._apply_scan_effect(pdf_data, client.id)

            response = HttpResponse(pdf_data, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{property_obj.name}_contrato_firmado.pdf"'

            os.remove(temp_doc_path)
            if os.path.exists(temp_pdf_path):
                os.remove(temp_pdf_path)

            return response

        except Clients.DoesNotExist:
            return Response({'error': 'Cliente no encontrado'}, status=404)
        except Property.DoesNotExist:
            return Response({'error': 'Propiedad no encontrada'}, status=404)
        except ValueError as e:
            return Response({'error': str(e)}, status=400)
        except subprocess.CalledProcessError:
            return Response({'error': 'Error al convertir DOCX a PDF'}, status=500)
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f'contrato_firma error: {e}')
            return Response({'error': str(e)}, status=400)

    def _get_firma_bytes(self, client):
        """
        Obtiene la firma digitalizada del cliente desde RENIEC/cache,
        quita el fondo blanco y retorna los bytes PNG.
        Solo funciona para clientes con DNI peruano.
        """
        import base64
        from PIL import Image

        if client.document_type != 'dni':
            return None

        from apps.reniec.models import DNICache
        try:
            cache = DNICache.objects.get(dni=client.number_doc)
            firma_b64 = cache.firma
        except DNICache.DoesNotExist:
            try:
                from apps.reniec.service import ReniecService
                success, result = ReniecService.lookup(
                    dni=client.number_doc,
                    source_app='contrato_firma',
                    include_photo=True,
                )
                if success:
                    cache = DNICache.objects.filter(dni=client.number_doc).first()
                    firma_b64 = cache.firma if cache else None
                else:
                    firma_b64 = None
            except Exception:
                firma_b64 = None

        if not firma_b64:
            return None

        try:
            img_data = base64.b64decode(firma_b64)
            img = Image.open(io.BytesIO(img_data)).convert("RGBA")

            # Quitar fondo blanco
            pixels = list(img.getdata())
            new_pixels = []
            for p in pixels:
                if p[0] > 230 and p[1] > 230 and p[2] > 230:
                    new_pixels.append((255, 255, 255, 0))
                else:
                    new_pixels.append(p)
            img.putdata(new_pixels)

            # Adelgazar trazos si la firma es muy gruesa
            img = self._normalize_firma_thickness(img)

            buffer = io.BytesIO()
            img.save(buffer, format="PNG")
            return buffer.getvalue()
        except Exception:
            return None

    def _normalize_firma_thickness(self, img):
        """
        Detecta firmas con trazos muy gruesos (lapicero grueso) y los adelgaza.
        Mide la densidad de píxeles de tinta dentro del bounding box de la firma.
        """
        import numpy as np
        from PIL import ImageFilter

        alpha = np.array(img.split()[3])

        # Bounding box de la firma (píxeles opacos)
        rows = np.any(alpha > 128, axis=1)
        cols = np.any(alpha > 128, axis=0)
        if not rows.any() or not cols.any():
            return img

        rmin, rmax = np.where(rows)[0][[0, -1]]
        cmin, cmax = np.where(cols)[0][[0, -1]]

        # Densidad = proporción de píxeles de tinta en el bounding box
        bbox_alpha = alpha[rmin:rmax + 1, cmin:cmax + 1]
        density = np.sum(bbox_alpha > 128) / bbox_alpha.size

        # Firma normal: ~5-15% densidad. Gruesa: >18%
        if density > 0.18:
            alpha_img = img.split()[3]

            # Más pasadas = más adelgazamiento
            if density > 0.35:
                passes = 3
            elif density > 0.25:
                passes = 2
            else:
                passes = 1

            for _ in range(passes):
                alpha_img = alpha_img.filter(ImageFilter.MinFilter(size=3))

            r, g, b, _ = img.split()
            img = Image.merge('RGBA', (r, g, b, alpha_img))

        return img

    def _insert_firma_anchor(self, docx_path, firma_png_bytes):
        """
        Abre el .docx renderizado, elimina el párrafo vacío de {{ firma }},
        e inserta la firma como imagen flotante (wp:anchor) encima de la línea ____.
        Esto iguala la estructura de ambas columnas para que las líneas queden a la misma altura.
        """
        from docx import Document
        from docx.shared import Cm
        from docx.oxml import OxmlElement
        from PIL import Image

        doc = Document(docx_path)

        if not doc.tables:
            return
        table = doc.tables[-1]
        if len(table.rows[0].cells) < 2:
            return
        right_cell = table.rows[0].cells[1]

        # Encontrar el párrafo con ____ y eliminar el párrafo vacío anterior
        cell_paragraphs = list(right_cell.paragraphs)
        underline_idx = None
        for i, p in enumerate(cell_paragraphs):
            if '____' in p.text:
                underline_idx = i
                break

        if underline_idx is None:
            return

        # Eliminar párrafo vacío anterior (era {{ firma }} renderizado como '')
        if underline_idx > 0:
            prev = cell_paragraphs[underline_idx - 1]
            if prev.text.strip() == '':
                prev._element.getparent().remove(prev._element)

        # Re-encontrar el párrafo ____
        underline_p = None
        for p in right_cell.paragraphs:
            if '____' in p.text:
                underline_p = p
                break

        if underline_p is None:
            doc.save(docx_path)
            return

        # Calcular dimensiones de la imagen
        img = Image.open(io.BytesIO(firma_png_bytes))
        target_width_emu = int(7 * 914400 / 2.54)  # 7cm en EMU
        aspect = img.height / img.width
        target_height_emu = int(target_width_emu * aspect)

        # Agregar imagen inline vía python-docx (gestiona el empaquetado)
        run = underline_p.add_run()
        run.add_picture(io.BytesIO(firma_png_bytes), width=Cm(7))

        # Encontrar el wp:inline y convertirlo a wp:anchor (flotante)
        from docx.oxml.ns import qn
        drawing = run._element.find(qn('w:drawing'))
        inline = drawing.find(qn('wp:inline'))

        if inline is None:
            doc.save(docx_path)
            return

        cx = inline.find(qn('wp:extent')).get('cx')
        cy = inline.find(qn('wp:extent')).get('cy')

        # Extraer hijos que necesitamos del inline
        docPr = inline.find(qn('wp:docPr'))
        cNvGfp = inline.find(qn('wp:cNvGraphicFramePr'))
        graphic = inline.find(qn('a:graphic'))

        # Construir wp:anchor
        anchor = OxmlElement('wp:anchor')
        anchor.set('distT', '0')
        anchor.set('distB', '0')
        anchor.set('distL', '114300')
        anchor.set('distR', '114300')
        anchor.set('simplePos', '0')
        anchor.set('relativeHeight', '251660288')
        anchor.set('behindDoc', '0')
        anchor.set('locked', '0')
        anchor.set('layoutInCell', '1')
        anchor.set('allowOverlap', '1')

        # simplePos
        spos = OxmlElement('wp:simplePos')
        spos.set('x', '0')
        spos.set('y', '0')
        anchor.append(spos)

        # Posición horizontal — centrado en la celda
        posH = OxmlElement('wp:positionH')
        posH.set('relativeFrom', 'column')
        offH = OxmlElement('wp:posOffset')
        offH.text = '100000'  # ~1mm desde el borde izquierdo de la celda
        posH.append(offH)
        anchor.append(posH)

        # Posición vertical — encima de la línea ____
        posV = OxmlElement('wp:positionV')
        posV.set('relativeFrom', 'paragraph')
        offV = OxmlElement('wp:posOffset')
        offV.text = str(-int(int(cy) * 0.55))  # La firma cruza la línea — 55% arriba, 45% sobre/debajo
        posV.append(offV)
        anchor.append(posV)

        # extent
        ext = OxmlElement('wp:extent')
        ext.set('cx', cx)
        ext.set('cy', cy)
        anchor.append(ext)

        # effectExtent
        ee = OxmlElement('wp:effectExtent')
        ee.set('l', '0')
        ee.set('t', '0')
        ee.set('r', '0')
        ee.set('b', '0')
        anchor.append(ee)

        # wrapNone — la imagen flota sobre el texto
        anchor.append(OxmlElement('wp:wrapNone'))

        # docPr, cNvGraphicFramePr, graphic (del inline original)
        anchor.append(docPr)
        anchor.append(cNvGfp)
        anchor.append(graphic)

        # Reemplazar inline con anchor
        drawing.remove(inline)
        drawing.append(anchor)

        doc.save(docx_path)

    def _apply_scan_effect(self, pdf_data, client_id):
        """
        Aplica un efecto de documento escaneado al PDF.
        Determinístico por cliente: mismo client_id = mismo efecto siempre.
        Diferentes clientes = diferentes escaneos.
        """
        import random
        import fitz  # PyMuPDF
        import numpy as np
        from PIL import Image, ImageFilter, ImageEnhance

        # Seed basado en client_id para reproducibilidad por cliente
        seed = hash(str(client_id)) % (2**32)
        rng = random.Random(seed)
        np_rng = np.random.RandomState(seed)

        pdf_doc = fitz.open(stream=pdf_data, filetype="pdf")
        scanned_pages = []

        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]

            # Renderizar a imagen (DPI ligeramente variable como escáner real)
            dpi = rng.uniform(245, 255)
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # 1. Rotación leve (papel no perfectamente alineado)
            angle = rng.uniform(-0.8, 0.8)
            fill_color = (rng.randint(250, 254), rng.randint(249, 253), rng.randint(247, 251))
            img = img.rotate(angle, resample=Image.BICUBIC, expand=False, fillcolor=fill_color)

            # 2. Tono cálido sutil (escáneres tienden a dar tono ligeramente amarillento)
            img_array = np.array(img, dtype=np.float32)
            img_array[:, :, 0] = np.clip(img_array[:, :, 0] + rng.uniform(0, 2.5), 0, 255)
            img_array[:, :, 1] = np.clip(img_array[:, :, 1] + rng.uniform(-1, 1), 0, 255)
            img_array[:, :, 2] = np.clip(img_array[:, :, 2] + rng.uniform(-3, -0.5), 0, 255)

            # 3. Ruido sutil (grano de escáner) — seed fijo por cliente
            noise_level = rng.uniform(1.8, 3.5)
            noise = np_rng.normal(0, noise_level, img_array.shape).astype(np.float32)
            img_array = np.clip(img_array + noise, 0, 255)
            img = Image.fromarray(img_array.astype(np.uint8))

            # 4. Ajuste de contraste y brillo
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(rng.uniform(0.96, 1.03))
            enhancer = ImageEnhance.Brightness(img)
            img = enhancer.enhance(rng.uniform(0.98, 1.02))

            # 5. Desenfoque muy leve (óptica del escáner)
            img = img.filter(ImageFilter.GaussianBlur(radius=rng.uniform(0.2, 0.45)))

            # 6. Compresión JPEG (artefactos típicos de escáner)
            jpeg_quality = rng.randint(88, 93)
            jpeg_buffer = io.BytesIO()
            img.save(jpeg_buffer, format="JPEG", quality=jpeg_quality)
            jpeg_buffer.seek(0)
            img = Image.open(jpeg_buffer).convert("RGB")

            scanned_pages.append(img)

        pdf_doc.close()

        # Guardar como PDF multi-página
        output = io.BytesIO()
        if len(scanned_pages) == 1:
            scanned_pages[0].save(output, format="PDF", resolution=250)
        else:
            scanned_pages[0].save(
                output, format="PDF", resolution=250,
                save_all=True, append_images=scanned_pages[1:]
            )

        return output.getvalue()

class DeleteRecipeApiView(generics.DestroyAPIView):
    queryset = RentalReceipt.objects.all()
    serializer_class = ReciptSerializer

    def get_queryset(self):
        queryset = super().get_queryset()
        return queryset.filter(reservation__seller=self.request.user)


class GetICSApiView(APIView):
    serializer_class = None
    permission_classes = [AllowAny]

    def get(self, request):

        if self.request.query_params:
            if self.request.query_params.get('q'):
                casa_sluged_name = slugify(self.request.query_params.get('q'))

                directory = str(Path(__file__).parent.parent.parent) + "/media/"
                f = open(os.path.join(directory, f'{casa_sluged_name}.ics'), 'rb')

                # Crear una respuesta HTTP con el contenido del archivo .ics
                response = HttpResponse(f, content_type='text/calendar')

                # Establecer el encabezado de Content-Disposition para descargar el archivo
                response['Content-Disposition'] = 'attachment; filename="evento.ics"'

        return response


class UpdateICSApiView(APIView):
    serializer_class = None

    def get(self, request):
        confeccion_ics()
        return Response({'message': 'ok'}, status=200)


class ProfitApiView(APIView):
    serializer_class = None

    def get(self, request):
        rta = {}

        evaluate_year = int(self.request.query_params['year'])
        for m in range(1, 13):
            last_day_month = calendar.monthrange(evaluate_year, m)[1]

            _, _, _, _, _, total_facturado = get_stadistics_period(
                datetime(evaluate_year, m, 1),
                last_day_month
            )

            rta[get_month_name(m)] = total_facturado

        return Response(
            rta,
            status=200
        )


############ Austin MOD ############
class VistaCalendarioApiView(viewsets.ModelViewSet):
    serializer_class = CalendarReservationSerializer # Usar el nuevo serializer ligero
    queryset = Reservation.objects.exclude(deleted=True).order_by("check_in_date")
    search_fields = [
        "client__email",
        "client__first_name",
        "client__last_name",
        "property__name"
    ]
    permission_classes = [IsAuthenticated]  # Requiere autenticación
    pagination_class = None  # Desactiva la paginación

    def get_queryset(self):
        queryset = super().get_queryset()

        if self.action == 'list':
            if self.request.query_params:
                from_param = self.request.query_params.get('from')
                type_param = self.request.query_params.get('type')
                now = datetime.now()
                check_in_time = time(15, 0)  # 3 PM
                check_out_time = time(11, 0)  # 11 AM

                # Modificación: Manejo del parámetro 'year' sin depender de 'month'
                year_param = self.request.query_params.get('year')
                month_param = self.request.query_params.get('month')

                if year_param:
                    try:
                        year_param = int(year_param)
                        if year_param < 1:
                            raise ValidationError({"error": "El parámetro 'year' debe ser un número entero positivo"})
                    except ValueError:
                        raise ValidationError({"error": "El parámetro 'year' debe ser un número válido"})

                    if month_param:
                        # Si 'month' también está presente
                        try:
                            month_param = int(month_param)
                            if month_param not in range(1, 13):
                                raise ValidationError({"error": "El parámetro 'month' debe ser un número entre 1 y 12"})
                        except ValueError:
                            raise ValidationError({"error": "El parámetro 'month' debe ser un número válido"})

                        # Filtrar por año y mes específico
                        last_day_month = calendar.monthrange(year_param, month_param)[1]
                        range_evaluate = (datetime(year_param, month_param, 1), datetime(year_param, month_param, last_day_month))

                        queryset = queryset.filter(
                            Q(check_in_date__range=range_evaluate) |
                            Q(check_out_date__range=range_evaluate)
                        )
                    else:
                        # Solo filtrar por año completo
                        start_of_year = datetime(year_param, 1, 1)
                        end_of_year = datetime(year_param, 12, 31)
                        queryset = queryset.filter(
                            Q(check_in_date__range=(start_of_year, end_of_year)) |
                            Q(check_out_date__range=(start_of_year, end_of_year))
                        )

                # Aquí se debe agregar el filtro para 'pending'
                if from_param == 'pending':
                    queryset = queryset.filter(status='pending')

                # Otros filtros existentes
                if from_param == 'today':
                    queryset = queryset.filter(check_in_date__gte=now)
                elif from_param == 'in_progress':
                    now_time = now.time()
                    today_date = now.date()
                    tomorrow_date = today_date + timedelta(days=1)

                    queryset = queryset.filter(
                        Q(check_in_date__lt=today_date, check_out_date__gt=today_date) |
                        (Q(check_in_date=today_date) & Q(check_out_date__gt=today_date)) |
                        (Q(check_out_date=today_date) & Q(check_out_date__gt=now)) |
                        (Q(check_out_date=tomorrow_date) & Q(check_out_date__gt=datetime.combine(tomorrow_date, check_out_time)))
                    )

                if type_param:
                    queryset = queryset.filter(origin=type_param)

        elif self.action in ['partial_update', 'update', 'destroy']:
            if not check_user_has_rol("admin", self.request.user):
                queryset = queryset.filter(
                    Q(origin="air") |
                    Q(seller=self.request.user)
                )

        if self.request.query_params.get('exclude'):
            queryset = queryset.exclude(origin=self.request.query_params['exclude'])

        return queryset.exclude(deleted=True)

    def get_serializer_class(self):
        if self.action == 'retrieve':
            return ReservationRetrieveSerializer
        if self.action == 'list':
            return CalendarReservationSerializer

        return super().get_serializer_class()

    def get_serializer_context(self):
        context = super().get_serializer_context()
        if self.action == 'retrieve':
            context['retrieve'] = True
        return context

    @extend_schema(
        parameters=[
            OpenApiParameter(
                "year",
                OpenApiTypes.INT,
                required=False,
                description="Filter results by year",
                enum=[2024, 2023, 2022],
            ),
            OpenApiParameter(
                "month",
                OpenApiTypes.INT,
                required=False,
                description="Filter results by month (number 1 to 12)",
                enum=list(range(1, 13)),
            ),
            OpenApiParameter(
                "from",
                OpenApiTypes.STR,
                description="Obtiene todas las reservas desde la fecha de hoy en adelante. Se puede combinar con type (tipo de reserva), pero no con year o month",
                required=False,
                enum=["today", "in_progress", "pending"]
            ),
            OpenApiParameter(
                "from_check_in",
                OpenApiTypes.STR,
                description="Obtiene todas las reservas que hayan comenzado este mes. Se usa en combinacion de los queryparams year y month",
                required=False,
                enum=["true"]
            ),
            OpenApiParameter(
                "type",
                OpenApiTypes.STR,
                description="Filtra las resevas según donde se generaron, AriBnB (air), Sistema Casa Austin (aus), Mantenimiento (man)",
                required=False,
                enum=["aus", "air", "man"]
            ),
            OpenApiParameter(
                "exclude",
                OpenApiTypes.STR,
                description="Excluye las resevas según donde se generaron, AriBnB (air), Sistema Casa Austin (aus), Mantenimiento (man)",
                required=False,
                enum=["aus", "air", "man"]
            ),
        ],
        responses={200: ReservationListSerializer(many=True)},
        methods=["GET"],
    )
    def list(self, request, *args, **kwargs):
        queryset = self.get_queryset()
        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)

    def perform_create(self, serializer):
        with transaction.atomic():
            user_seller = self.request.user
            if self.request.POST['origin'].lower() == 'air':
                user_seller = CustomUser.objects.get(first_name='AirBnB')

            # Para reservas de cliente, establecer status pending y deadline
            status = 'approved'
            if self.request.POST.get('origin', '').lower() == 'client':
                status = 'pending'

            instance = serializer.save(seller=user_seller, status=status)

            # Si es reserva de cliente, establecer deadline de 1 hora
            if instance.origin == 'client':
                from django.utils import timezone
                instance.payment_voucher_deadline = timezone.now() + timedelta(hours=1)
                instance.save()
            if instance.late_checkout:
                original_check_out_date = instance.check_out_date - timedelta(days=1)
                instance.late_check_out_date = original_check_out_date
                instance.check_out_date = original_check_out_date + timedelta(days=1)
                instance.save()

            for file in self.request.FILES.getlist('file'):
                RentalReceipt.objects.create(
                    reservation=instance,
                    file=file
                )

        confeccion_ics()

        generate_audit(
            instance,
            self.request.user,
            "create",
            "Reserva creada"
        )

    def perform_update(self, serializer):
        with transaction.atomic():
            instance = serializer.save()
            if instance.late_checkout:
                if instance.late_check_out_date is None:
                    original_check_out_date = instance.check_out_date - timedelta(days=1)
                    instance.late_check_out_date = original_check_out_date
                    instance.check_out_date = original_check_out_date + timedelta(days=1)
                instance.save()

            confeccion_ics()

            generate_audit(
                instance,
                self.request.user,
                "update",
                "Reserva actualizada full"
            )

        return super().perform_update(serializer)

    def partial_update(self, request, *args, **kwargs):
        instance = self.get_object()
        serializer = self.get_serializer(instance, data=request.data, partial=True)
        serializer.is_valid(raise_exception=True)

        with transaction.atomic():
            instance = serializer.save()
            if instance.late_checkout:
                if instance.late_check_out_date is None:
                    original_check_out_date = instance.check_out_date - timedelta(days=1)
                    instance.late_check_out_date = original_check_out_date
                    instance.check_out_date = original_check_out_date + timedelta(days=1)
                instance.save()

            for file in request.FILES.getlist('file'):
                RentalReceipt.objects.create(
                    reservation=instance,
                    file=file
                )

        confeccion_ics()

        generate_audit(
            serializer.instance,
            self.request.user,
            "update",
            "Reserva actualizada"
        )
        return Response(serializer.data)


    def destroy(self, request, *args, **kwargs):
        instance = self.get_object()
        # Pasar el motivo de eliminación para auditoría con simple_history
        user_info = f"{request.user.email}" if request.user and hasattr(request.user, 'email') else "Usuario API"
        instance.delete(reason=f"API Calendario: eliminado por {user_info}")

        confeccion_ics()

        generate_audit(
            instance,
            self.request.user,
            "delete",
            "Reserva eliminada"
        )
        return Response(status=204)

@csrf_exempt
def confirm_reservation(request, uuid):
    if request.method != "POST":
        return HttpResponseBadRequest("Método no permitido")

    # Buscar reserva por ID o uuid_external (sin guiones)
    try:
        reservation = Reservation.objects.get(id=uuid)
    except Reservation.DoesNotExist:
        uuid_clean = uuid.replace("-", "")
        reservation = get_object_or_404(Reservation, uuid_external=uuid_clean)

    # Intentar leer JSON, sino usar POST clásico
    try:
        data = json.loads(request.body)
    except (json.JSONDecodeError, UnicodeDecodeError):
        data = request.POST

    # Obtener IP real del cliente (considerando proxy)
    ip = request.META.get("HTTP_X_FORWARDED_FOR")
    if ip:
        ip = ip.split(",")[0]
    else:
        ip = request.META.get("REMOTE_ADDR")

    # Guardar datos de navegación
    reservation.ip_cliente = ip
    reservation.user_agent = request.META.get("HTTP_USER_AGENT", "")
    reservation.referer = request.META.get("HTTP_REFERER", "")
    reservation.fbclid = data.get("fbclid", "")
    reservation.utm_source = data.get("utm_source", "")
    reservation.utm_medium = data.get("utm_medium", "")
    reservation.utm_campaign = data.get("utm_campaign", "")
    reservation.fbc = data.get("fbc", "")
    reservation.fbp = data.get("fbp", "")
    reservation.save()

    # Enviar a Meta
    send_purchase_event_to_meta(
        phone=reservation.client.tel_number,
        email=reservation.client.email,
        first_name=reservation.client.first_name,
        last_name=reservation.client.last_name,
        amount=reservation.price_usd,
        currency="USD",
        ip=reservation.ip_cliente,
        user_agent=reservation.user_agent,
        fbc=reservation.fbc,
        fbp=reservation.fbp,
        fbclid=reservation.fbclid,
        utm_source=reservation.utm_source,
        utm_medium=reservation.utm_medium,
        utm_campaign=reservation.utm_campaign,
        birthday=str(reservation.client.date) if reservation.client.date else None  # <-- aquí

    )

    return JsonResponse({"message": "✅ ¡Reserva confirmada correctamente!"})


class MonthlyReservationsExportAPIView(APIView):
    """
    Endpoint para exportar datos de reservas por mes
    """
    permission_classes = [IsAuthenticated]  # Requiere autenticación

    @extend_schema(
        parameters=[
            OpenApiParameter(
                "year",
                OpenApiTypes.INT,
                required=True,
                description="Año de las reservas a exportar",
                location=OpenApiParameter.QUERY
            ),
            OpenApiParameter(
                "month",
                OpenApiTypes.INT,
                required=True,
                description="Mes de las reservas a exportar (1-12)",
                location=OpenApiParameter.QUERY
            ),
        ],
        responses={
            200: {
                "type": "object",
                "properties": {
                    "success": {"type": "boolean"},
                    "data": {
                        "type": "object",
                        "properties": {
                            "period": {"type": "string"},
                            "total_reservations": {"type": "integer"},
                            "reservations": {
                                "type": "array",
                                "items": {
                                    "type": "object",
                                    "properties": {
                                        "id": {"type": "integer"},
                                        "client_name": {"type": "string"},
                                        "client_email": {"type": "string"},
                                        "client_phone": {"type": "string"},
                                        "property_name": {"type": "string"},
                                        "check_in_date": {"type": "string", "format": "date"},
                                        "check_out_date": {"type": "string", "format": "date"},
                                        "guests": {"type": "integer"},
                                        "price_usd": {"type": "number"},
                                        "price_sol": {"type": "number"},
                                        "advance_payment": {"type": "number"},
                                        "advance_payment_currency": {"type": "string"},
                                        "full_payment": {"type": "boolean"},
                                        "temperature_pool": {"type": "boolean"},
                                        "origin": {"type": "string"},
                                        "status": {"type": "string"},
                                        "seller_name": {"type": "string"},
                                        "created": {"type": "string", "format": "date-time"},
                                        "number_nights": {"type": "integer"},
                                        "points_redeemed": {"type": "number"},
                                        "discount_code_used": {"type": "string"}
                                    }
                                }
                            }
                        }
                    }
                }
            },
            400: {
                "type": "object",
                "properties": {
                    "success": {"type": "boolean"},
                    "error": {"type": "string"}
                }
            }
        }
    )
    def get(self, request):
        try:
            # Validar parámetros requeridos
            year_param = request.query_params.get('year')
            month_param = request.query_params.get('month')

            if not year_param or not month_param:
                return Response({
                    "success": False,
                    "error": "Los parámetros 'year' y 'month' son requeridos"
                }, status=400)

            # Validar año
            try:
                year = int(year_param)
                if year < 2020 or year > 2030:
                    raise ValueError()
            except ValueError:
                return Response({
                    "success": False,
                    "error": "El parámetro 'year' debe ser un número válido entre 2020 y 2030"
                }, status=400)

            # Validar mes
            try:
                month = int(month_param)
                if month < 1 or month > 12:
                    raise ValueError()
            except ValueError:
                return Response({
                    "success": False,
                    "error": "El parámetro 'month' debe ser un número entre 1 y 12"
                }, status=400)

            # Calcular rango de fechas para el mes
            last_day_month = calendar.monthrange(year, month)[1]
            start_date = datetime(year, month, 1).date()
            end_date = datetime(year, month, last_day_month).date()

            # Obtener reservas del mes
            reservations = Reservation.objects.filter(
                check_in_date__gte=start_date,
                check_in_date__lte=end_date,
                deleted=False
            ).select_related('client', 'property', 'seller').order_by('check_in_date')

            # Formatear datos de reservas
            reservations_data = []
            for reservation in reservations:
                # Calcular número de noches
                if reservation.check_in_date and reservation.check_out_date:
                    delta = reservation.check_out_date - reservation.check_in_date
                    number_nights = delta.days
                else:
                    number_nights = 0

                # Formatear nombre del cliente
                client_name = ""
                client_email = ""
                client_phone = ""
                if reservation.client:
                    first_name = reservation.client.first_name or ""
                    last_name = reservation.client.last_name or ""
                    client_name = f"{first_name} {last_name}".strip()
                    client_email = reservation.client.email or ""
                    client_phone = reservation.client.tel_number or ""

                # Formatear nombre del vendedor
                seller_name = ""
                if reservation.seller:
                    seller_first = reservation.seller.first_name or ""
                    seller_last = reservation.seller.last_name or ""
                    seller_name = f"{seller_first} {seller_last}".strip()

                # Mapear status para mejor legibilidad
                status_mapping = {
                    'approved': 'Aprobada',
                    'pending': 'Pendiente',
                    'incomplete': 'Incompleta',
                    'rejected': 'Rechazada',
                    'cancelled': 'Cancelada'
                }
                status_display = status_mapping.get(reservation.status, reservation.status)

                # Mapear origen para mejor legibilidad
                origin_mapping = {
                    'air': 'Airbnb',
                    'aus': 'Austin',
                    'man': 'Mantenimiento',
                    'client': 'Cliente Web'
                }
                origin_display = origin_mapping.get(reservation.origin, reservation.origin)

                reservations_data.append({
                    "id": reservation.id,
                    "client_name": client_name,
                    "client_email": client_email,
                    "client_phone": client_phone,
                    "property_name": reservation.property.name if reservation.property else "",
                    "check_in_date": reservation.check_in_date.strftime('%Y-%m-%d'),
                    "check_out_date": reservation.check_out_date.strftime('%Y-%m-%d'),
                    "guests": reservation.guests,
                    "price_usd": float(reservation.price_usd or 0),
                    "price_sol": float(reservation.price_sol or 0),
                    "advance_payment": float(reservation.advance_payment or 0),
                    "advance_payment_currency": reservation.advance_payment_currency,
                    "full_payment": reservation.full_payment,
                    "temperature_pool": reservation.temperature_pool,
                    "origin": origin_display,
                    "status": status_display,
                    "seller_name": seller_name,
                    "created": reservation.created.isoformat() if reservation.created else "",
                    "number_nights": number_nights,
                    "points_redeemed": float(reservation.points_redeemed or 0),
                    "discount_code_used": reservation.discount_code_used or "",
                    "tel_contact_number": reservation.tel_contact_number or "",
                    "comentarios_reservas": reservation.comentarios_reservas or ""
                })

            # Formatear nombre del período
            month_names = {
                1: 'Enero', 2: 'Febrero', 3: 'Marzo', 4: 'Abril',
                5: 'Mayo', 6: 'Junio', 7: 'Julio', 8: 'Agosto',
                9: 'Septiembre', 10: 'Octubre', 11: 'Noviembre', 12: 'Diciembre'
            }
            period_name = f"{month_names[month]} {year}"

            return Response({
                "success": True,
                "data": {
                    "period": period_name,
                    "total_reservations": len(reservations_data),
                    "reservations": reservations_data
                }
            })

        except Exception as e:
            return Response({
                "success": False,
                "error": f"Error interno del servidor: {str(e)}"
            }, status=500)


class PropertyCalendarOccupancyAPIView(APIView):
    """
    Endpoint para obtener la ocupación del calendario de una propiedad
    """
    permission_classes = [AllowAny]

    @extend_schema(
        parameters=[
            OpenApiParameter(
                "year",
                OpenApiTypes.INT,
                required=True,
                description="Año para filtrar las reservas",
                location=OpenApiParameter.QUERY
            ),
            OpenApiParameter(
                "month",
                OpenApiTypes.INT,
                required=False,
                description="Mes para filtrar las reservas (1-12). Si no se envía, devuelve todo el año",
                location=OpenApiParameter.QUERY
            ),
            OpenApiParameter(
                "day",
                OpenApiTypes.INT,
                required=False,
                description="Día para filtrar las reservas (1-31). Requiere que se especifique el mes",
                location=OpenApiParameter.QUERY
            ),
        ],
        responses={
            200: {
                "type": "object",
                "properties": {
                    "success": {"type": "boolean"},
                    "data": {
                        "type": "object",
                        "properties": {
                            "property_id": {"type": "string"},
                            "occupancy": {
                                "type": "array",
                                "items": {
                                    "type": "object",
                                    "properties": {
                                        "start_date": {"type": "string", "format": "date"},
                                        "end_date": {"type": "string", "format": "date"},
                                        "guest_name": {"type": "string"},
                                        "status": {"type": "string", "enum": ["confirmed", "pending", "incomplete"]},
                                        "reservation_id": {"type": "string"}
                                    }
                                }
                            }
                        }
                    }
                }
            },
            404: {
                "type": "object",
                "properties": {
                    "success": {"type": "boolean"},
                    "error": {"type": "string"}
                }
            }
        }
    )
    def get(self, request, property_id):
        try:
            # Buscar propiedad por ID o slug
            try:
                if property_id.isdigit():
                    property_obj = Property.objects.get(id=property_id, deleted=False)
                else:
                    property_obj = Property.objects.get(slug=property_id, deleted=False)
            except Property.DoesNotExist:
                return Response({
                    "success": False,
                    "error": "Propiedad no encontrada"
                }, status=404)

            # Validar parámetros
            year_param = request.query_params.get('year')
            month_param = request.query_params.get('month')
            day_param = request.query_params.get('day')

            if not year_param:
                return Response({
                    "success": False,
                    "error": "El parámetro 'year' es requerido"
                }, status=400)

            try:
                year = int(year_param)
                if year < 1:
                    raise ValueError()
            except ValueError:
                return Response({
                    "success": False,
                    "error": "El parámetro 'year' debe ser un número entero positivo"
                }, status=400)

            if month_param:
                try:
                    month = int(month_param)
                    if month < 1 or month > 12:
                        raise ValueError()
                except ValueError:
                    return Response({
                        "success": False,
                        "error": "El parámetro 'month' debe ser un número entre 1 y 12"
                    }, status=400)

            if day_param:
                if not month_param:
                    return Response({
                        "success": False,
                        "error": "El parámetro 'day' requiere que también se especifique 'month'"
                    }, status=400)

                try:
                    day = int(day_param)
                    if day < 1 or day > 31:
                        raise ValueError()

                    # Validar que el día sea válido para el mes y año especificados
                    max_day = calendar.monthrange(year, month)[1]
                    if day > max_day:
                        return Response({
                            "success": False,
                            "error": f"El día {day} no es válido para {month}/{year}. El mes tiene {max_day} días"
                        }, status=400)

                except ValueError:
                    return Response({
                        "success": False,
                        "error": "El parámetro 'day' debe ser un número entre 1 y 31"
                    }, status=400)

            # Construir filtro de fechas
            if day_param:
                # Filtrar por día específico
                start_date = datetime(year, month, day).date()
                end_date = datetime(year, month, day).date()
            elif month_param:
                # Filtrar por mes específico
                last_day_month = calendar.monthrange(year, month)[1]
                start_date = datetime(year, month, 1).date()
                end_date = datetime(year, month, last_day_month).date()
            else:
                # Filtrar por todo el año
                start_date = datetime(year, 1, 1).date()
                end_date = datetime(year, 12, 31).date()

            # Obtener reservas de la propiedad en el rango de fechas
            reservations = Reservation.objects.filter(
                property=property_obj,
                deleted=False,
                status__in=['approved', 'pending', 'incomplete', 'under_review']
            ).filter(
                Q(check_in_date__lte=end_date) & Q(check_out_date__gte=start_date)
            ).select_related('client').order_by('check_in_date')

            # Formatear datos de ocupación
            occupancy_data = []
            for reservation in reservations:
                # Formatear nombre del huésped
                if reservation.client:
                    first_name = reservation.client.first_name or ""
                    last_name = reservation.client.last_name or ""

                    # Crear formato con primer nombre e inicial del primer apellido
                    if first_name and last_name:
                        # Obtener solo el primer nombre y la inicial del primer apellido
                        primer_nombre = first_name.split()[0] if first_name else ""
                        guest_name = f"{primer_nombre} {last_name[0]}."
                    elif first_name:
                        primer_nombre = first_name.split()[0] if first_name else ""
                        guest_name = primer_nombre
                    elif last_name:
                        guest_name = f"{last_name[0]}."
                    else:
                        guest_name = "Cliente sin nombre"
                else:
                    guest_name = "Cliente sin datos"

                # Mapear status
                status_mapping = {
                    'approved': 'confirmed',
                    'pending': 'pending',
                    'incomplete': 'incomplete',
                    'rejected': 'rejected',
                    'cancelled': 'cancelled'
                }

                status = status_mapping.get(reservation.status, reservation.status)

                occupancy_data.append({
                    "start_date": reservation.check_in_date.strftime('%Y-%m-%d'),
                    "end_date": reservation.check_out_date.strftime('%Y-%m-%d'),
                    "guest_name": guest_name,
                    "status": status,
                    "reservation_id": str(reservation.id),
                    "origin": reservation.origin
                })

            return Response({
                "success": True,
                "data": {
                    "property_id": property_obj.slug or str(property_obj.id),
                    "occupancy": occupancy_data
                }
            })

        except Exception as e:
            return Response({
                "success": False,
                "error": f"Error interno del servidor: {str(e)}"
            }, status=500)


class QRReservationView(APIView):
    """
    Endpoint público para mostrar los datos de una reserva por su ID.
    Útil para códigos QR en recepción.
    """
    permission_classes = [AllowAny]
    
    def get(self, request, reservation_id):
        """
        Obtiene los datos de una reserva específica por su ID.
        """
        try:
            # Buscar la reserva por ID
            reservation = Reservation.objects.select_related('client', 'property').get(
                id=reservation_id,
                deleted=False
            )
            
            # Preparar datos del cliente
            client = reservation.client
            client_data = {
                "name": None,
                "facebook_photo": None,
                "facebook_link": False,
                "referral_code": None,
                "level": None,
                "level_icon": None,
                "referral_discount_percentage": None,
                "property_name": reservation.property.name
            }
            
            if client:
                # Nombre del cliente
                client_data["name"] = f"{client.first_name or ''} {client.last_name or ''}".strip() or "Cliente sin nombre"
                
                # Datos de Facebook
                if client.facebook_linked and client.facebook_profile_data:
                    client_data["facebook_photo"] = client.get_facebook_profile_picture()
                    client_data["facebook_link"] = True
                else:
                    client_data["facebook_link"] = False
                
                # Código de referido
                client_data["referral_code"] = client.get_referral_code() if hasattr(client, 'get_referral_code') else client.referral_code
                
                # Obtener nivel del cliente (achievement más alto)
                from apps.clients.models import ClientAchievement
                highest_achievement = ClientAchievement.objects.filter(
                    client=client,
                    deleted=False
                ).select_related('achievement').order_by(
                    '-achievement__required_reservations',
                    '-achievement__required_referrals',
                    '-achievement__required_referral_reservations'
                ).first()
                
                if highest_achievement:
                    client_data["level"] = highest_achievement.achievement.name
                    client_data["level_icon"] = highest_achievement.achievement.icon
                    
                    # Obtener descuento por referido de este nivel
                    from apps.property.models import ReferralDiscountByLevel
                    referral_discount = ReferralDiscountByLevel.objects.filter(
                        achievement=highest_achievement.achievement,
                        is_active=True,
                        deleted=False
                    ).first()
                    
                    if referral_discount:
                        client_data["referral_discount_percentage"] = float(referral_discount.discount_percentage)
            
            # Agregar información de check-in, check-out y late checkout
            client_data["check_in_date"] = reservation.check_in_date.isoformat()
            client_data["check_out_date"] = reservation.check_out_date.isoformat()
            client_data["late_checkout"] = reservation.late_checkout
            
            return Response({
                "success": True,
                "data": client_data
            })
            
        except Reservation.DoesNotExist:
            return Response({
                "success": False,
                "error": "Reserva no encontrada"
            }, status=404)
        except Exception as e:
            return Response({
                "success": False,
                "error": f"Error interno: {str(e)}"
            }, status=500)


class ActiveReservationsView(APIView):
    """
    GET /api/v1/active/
    Devuelve todas las reservas activas en este momento.
    Valida horarios de check-in (12 PM) y check-out (11 AM) en horario de Perú.

    Requiere autenticación JWT o token de portal (X-Portal-Key header).
    Variable de entorno: WIFI_PORTAL_SECRET_KEY
    """
    permission_classes = [AllowAny]  # Validación manual para soportar ambos métodos

    def get(self, request):
        # Verificar autenticación: JWT o Portal Key
        portal_key = request.headers.get('X-Portal-Key')
        secret_key = os.environ.get('WIFI_PORTAL_SECRET_KEY', '564y5r4564gergRTHRthdrghdFH')

        # DEBUG: Log para diagnosticar problema de autenticación
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f"ActiveReservationsView: portal_key={portal_key}, secret_key={secret_key}, user={request.user}, is_auth={getattr(request.user, 'is_authenticated', 'NO_ATTR')}")

        if not request.user.is_authenticated:
            # Si no hay JWT, verificar Portal Key
            if portal_key != secret_key:
                return Response(
                    {'error': 'Autenticación requerida'},
                    status=status.HTTP_401_UNAUTHORIZED
                )
        try:
            from django.utils import timezone
            
            # Obtener hora actual en horario de Perú
            local_now = timezone.localtime(timezone.now())
            now_date = local_now.date()
            now_time = local_now.time()
            
            checkin_time = time(12, 0)  # 12 PM (mediodía)
            checkout_time = time(11, 0)  # 11 AM
            
            # Obtener todas las reservas aprobadas que podrían estar activas
            reservations = Reservation.objects.filter(
                deleted=False,
                status='approved',
                check_in_date__lte=now_date,  # Ya empezó o empieza hoy
                check_out_date__gte=now_date   # No ha terminado o termina hoy
            ).select_related('property', 'client')
            
            active_reservations = []
            
            for res in reservations:
                # Verificar si la reserva está activa en este momento
                is_active = True
                
                # Si es el día de check-in, debe ser después de las 12 PM
                if now_date == res.check_in_date and now_time < checkin_time:
                    is_active = False
                
                # Si es el día de check-out, debe ser antes de las 11 AM
                if now_date == res.check_out_date and now_time >= checkout_time:
                    is_active = False
                
                # Si está fuera del rango de fechas
                if now_date < res.check_in_date or now_date > res.check_out_date:
                    is_active = False
                
                if is_active:
                    # Manejar reservas con y sin cliente (Airbnb, Mantenimiento)
                    client = res.client
                    
                    if client:
                        client_name = f"{client.first_name or ''} {client.last_name or ''}".strip() or "Sin nombre"
                        phone = client.tel_number or ''
                        referral_code = client.get_referral_code() if hasattr(client, 'get_referral_code') else (client.referral_code or '')
                    else:
                        # Reservas sin cliente (Airbnb, Mantenimiento)
                        origin_display = res.get_origin_display() if res.origin else "Reserva"
                        client_name = f"Reserva {origin_display}"
                        phone = res.tel_contact_number or ''
                        referral_code = ''
                    
                    active_reservations.append({
                        'id': str(res.id),
                        'property': res.property.name if res.property else 'Sin propiedad',
                        'property_id': res.property.player_id if res.property else None,
                        'property_color': res.property.background_color if res.property else '#2196F3',
                        'client_name': client_name,
                        'phone': phone,
                        'referral_code': referral_code,
                        'check_in_date': res.check_in_date.isoformat(),
                        'check_out_date': res.check_out_date.isoformat(),
                        'guests': res.guests,
                        'temperature_pool': res.temperature_pool,
                        'late_checkout': res.late_checkout,
                        'comentarios': res.comentarios_reservas or '',
                        'is_currently_active': True,
                        'origin': res.origin or ''
                    })
            
            # Obtener reservas que hacen check-in hoy
            checkin_today_reservations = Reservation.objects.filter(
                deleted=False,
                status='approved',
                check_in_date=now_date
            ).select_related('property', 'client')
            
            checkin_today = []
            
            for res in checkin_today_reservations:
                # Manejar reservas con y sin cliente
                client = res.client
                
                if client:
                    client_name = f"{client.first_name or ''} {client.last_name or ''}".strip() or "Sin nombre"
                    phone = client.tel_number or ''
                    referral_code = client.get_referral_code() if hasattr(client, 'get_referral_code') else (client.referral_code or '')
                else:
                    origin_display = res.get_origin_display() if res.origin else "Reserva"
                    client_name = f"Reserva {origin_display}"
                    phone = res.tel_contact_number or ''
                    referral_code = ''
                
                checkin_today.append({
                    'id': str(res.id),
                    'property': res.property.name if res.property else 'Sin propiedad',
                    'property_id': res.property.player_id if res.property else None,
                    'property_color': res.property.background_color if res.property else '#2196F3',
                    'client_name': client_name,
                    'phone': phone,
                    'referral_code': referral_code,
                    'check_in_date': res.check_in_date.isoformat(),
                    'check_out_date': res.check_out_date.isoformat(),
                    'guests': res.guests,
                    'temperature_pool': res.temperature_pool,
                    'late_checkout': res.late_checkout,
                    'comentarios': res.comentarios_reservas or '',
                    'checkin_time': '12:00 PM',
                    'origin': res.origin or ''
                })
            
            return Response({
                "success": True,
                "count": len(active_reservations),
                "active_reservations": active_reservations,
                "check_in_today": checkin_today
            })
            
        except Exception as e:
            return Response({
                "success": False,
                "error": f"Error interno: {str(e)}"
            }, status=500)